# Final_Demo.py (STREAMLIT) — Updated to compute real evaluation metrics locally
import streamlit as st
import requests
import datetime
from io import BytesIO
from docx import Document
import numpy as np
import re
import os
import html
from typing import Tuple, Optional

# Optional token encoder; if not installed fallback to simple estimate
try:
    import tiktoken
except Exception:
    tiktoken = None

# Optional semantic similarity using sentence-transformers (only used if installed;
# we will NOT attempt to load a huge model automatically).
try:
    from sentence_transformers import SentenceTransformer
    has_sentence_transformers = True
except Exception:
    SentenceTransformer = None
    has_sentence_transformers = False

# -----------------------------
# Backend URL (change for production)
# -----------------------------
BASE_URL = os.getenv("BASE_URL", "http://localhost:8000")

# -----------------------------
# Streamlit page setup
# -----------------------------
st.set_page_config(
    page_title="🌍 AI Trip Planner",
    layout="centered",
    initial_sidebar_state="expanded",
)
st.title("🌍 Demo for AI Tour Planner")

# -----------------------------
# Session State Initialization
# -----------------------------
if "plans" not in st.session_state:
    st.session_state.plans = []
if "trip_details_submitted" not in st.session_state:
    st.session_state.trip_details_submitted = False
if "departure_city" not in st.session_state:
    st.session_state.departure_city = ""
if "destination_city" not in st.session_state:
    st.session_state.destination_city = ""
if "departure_date" not in st.session_state:
    st.session_state.departure_date = None
if "return_date" not in st.session_state:
    st.session_state.return_date = None
if "relevance" not in st.session_state:
    st.session_state.relevance = 0.0
if "accuracy" not in st.session_state:
    st.session_state.accuracy = 0.0
if "tokens_in" not in st.session_state:
    st.session_state.tokens_in = 0
if "tokens_out" not in st.session_state:
    st.session_state.tokens_out = 0
if "total_tokens" not in st.session_state:
    st.session_state.total_tokens = 0
if "last_response" not in st.session_state:
    st.session_state.last_response = ""
if "semantic_model" not in st.session_state:
    st.session_state.semantic_model = None  # optional cached small model

# ----------------------------------------------------
# 🌐 TRANSLATOR TOOLBOX (SIDEBAR VERSION)
# ----------------------------------------------------
with st.sidebar:
    st.markdown("---")
    st.header("🌐 Language Barrier? Please use Sambhāṣhak")

    text_to_translate = st.text_area(
        "Enter text to translate (leave empty to translate entire itinerary):"
    )

    languages = {
        "Hindi": "hi",
        "Marathi": "mr",
        "Tamil": "ta",
        "Telugu": "te",
        "Kannada": "kn",
        "Gujarati": "gu",
        "Bengali": "bn",
        "Punjabi": "pa",
        "Malayalam": "ml",
        "Spanish": "es",
        "French": "fr",
        "German": "de",
        "Italian": "it",
        "Chinese": "zh",
        "Japanese": "ja",
    }

    selected_language = st.selectbox("Select your preferred language:", list(languages.keys()))
    lang_code = languages[selected_language]

    if st.button("Translate Text"):
        final_text = text_to_translate.strip() or st.session_state.last_response

        with st.spinner("Translating..."):
            try:
                payload = {"text": final_text, "target_language": lang_code}
                response = requests.post(f"{BASE_URL}/translate", json=payload, timeout=60)

                if response.status_code == 200:
                    result = response.json()
                    translated_text = result.get("translated_text", "")
                    audio_bytes = result.get("audio_bytes", None)

                    st.success(f"Translated to {selected_language}:")
                    st.write(translated_text)

                    st.session_state.plans.append((f"Tranxslated ({selected_language})", translated_text))

                    if audio_bytes:
                        st.audio(BytesIO(bytes(audio_bytes)), format="audio/mp3")
                    else:
                        st.info("Audio not available for this language.")

                else:
                    st.error(f"Translation error: {response.text}")

            except Exception as e:
                st.error(f"Translation server error: {e}")

# -----------------------------
# Token Count Helper
# -----------------------------
def estimate_tokens(text: str, model: str = "gpt-4o-mini"):
    if not text:
        return 0
    if tiktoken:
        try:
            # use generic encoding; fallback if specific model not present
            try:
                enc = tiktoken.encoding_for_model(model)
            except Exception:
                enc = tiktoken.get_encoding("cl100k_base")
            return len(enc.encode(text))
        except Exception:
            # fallback
            return int(len(text) / 4)
    # fallback rough estimate (1 token ≈ 4 chars)
    return int(len(text) / 4)

# -----------------------------
# Text helpers for evaluation
# -----------------------------
_stopwords = {
    "the","and","for","with","that","this","from","will","your","you","are",
    "a","an","to","in","on","of","it","is","as","by","at","or","be","we"
}

def normalize_text(s: str) -> str:
    if not s:
        return ""
    s = html.unescape(s)
    s = s.lower()
    # replace punctuation with spaces
    s = re.sub(r"[^\w\s\-\/]", " ", s)
    s = re.sub(r"\s+", " ", s).strip()
    return s

def tokenize(text: str):
    text = normalize_text(text)
    tokens = [t for t in text.split() if len(t) > 1 and t not in _stopwords]
    return tokens

def extract_keywords_from_query(query: str, max_keywords: int = 40):
    tokens = tokenize(query)
    # prefer longer tokens and unique
    tokens = sorted(set(tokens), key=lambda t: (-len(t), t))
    return tokens[:max_keywords]

# -----------------------------
# Optional semantic similarity
# -----------------------------
def get_semantic_similarity(a: str, b: str) -> Optional[float]:
    """
    Returns cosine similarity in [0,1] if sentence-transformers is available.
    If not available or loading fails, returns None so the code will fallback.
    """
    if not has_sentence_transformers:
        return None
    try:
        # load once and cache in session_state to avoid repeated heavy loads
        if st.session_state.semantic_model is None:
            # choose a small model (user machine must have it installed) to avoid heavy loads.
            # If it isn't present, this will raise and we'll fallback gracefully.
            st.session_state.semantic_model = SentenceTransformer("all-MiniLM-L6-v2")
        model = st.session_state.semantic_model
        emb_a = model.encode(a, convert_to_numpy=True)
        emb_b = model.encode(b, convert_to_numpy=True)
        # cosine similarity
        denom = np.linalg.norm(emb_a) * np.linalg.norm(emb_b)
        if denom == 0:
            return None
        sim = float(np.dot(emb_a, emb_b) / denom)
        # map [-1,1] -> [0,1]
        sim = max(-1.0, min(1.0, sim))
        sim01 = (sim + 1.0) / 2.0
        return sim01
    except Exception:
        # don't crash if model unavailable or any error; return None to fallback
        return None

# -----------------------------
# Compute evaluation metrics locally
# -----------------------------
def compute_evaluation(user_query: str, answer_text: str, departure_city: str, destination_city: str,
                       departure_date: Optional[datetime.date], return_date: Optional[datetime.date]) -> Tuple[float, float]:
    """
    Returns (relevance_score, accuracy_score) both in [0.0, 1.0].
    - relevance_score: mixture of token-overlap and optional semantic similarity
    - accuracy_score: fraction of required factual items present (departure/destination/dates)
    """
    # normalize
    q = normalize_text(user_query)
    a = normalize_text(answer_text)

    # 1) token-overlap based relevance
    q_tokens = set(extract_keywords_from_query(q, max_keywords=80))
    a_tokens = set(tokenize(a))
    if not q_tokens:
        overlap_ratio = 0.0
    else:
        common = q_tokens.intersection(a_tokens)
        overlap_ratio = len(common) / len(q_tokens)  # 0..1

    # 2) optional semantic similarity (if model available)
    sem_sim = get_semantic_similarity(user_query, answer_text)  # may be None

    # combine them: weight semantic sim higher if available
    if sem_sim is not None:
        relevance = 0.6 * sem_sim + 0.4 * overlap_ratio
    else:
        relevance = overlap_ratio

    # clamp
    relevance = float(min(1.0, max(0.0, relevance)))

    # -------------------------
    # accuracy: factual presence checks
    required_checks = []
    if departure_city and departure_city.strip():
        required_checks.append(("departure_city", normalize_text(departure_city)))
    if destination_city and destination_city.strip():
        required_checks.append(("destination_city", normalize_text(destination_city)))
    if departure_date:
        required_checks.append(("departure_date", departure_date.strftime("%Y-%m-%d")))
        # also accept human reading formats
        required_checks.append(("departure_date_alt", departure_date.strftime("%d-%b-%Y")))
    if return_date:
        required_checks.append(("return_date", return_date.strftime("%Y-%m-%d")))
        required_checks.append(("return_date_alt", return_date.strftime("%d-%b-%Y")))

    # Count how many required facts are found. For date checks allow either format to count as present.
    found = 0
    total_required = 0
    # we'll map logical facts to groups: departure city, destination city, dates (treated separately)
    # simpler: treat each unique "type" once, so we avoid double-counting alt date formats
    # Build groups
    groups = {}
    for key, value in required_checks:
        # group by prefix before underscore
        group = key.split("_")[0]
        groups.setdefault(group, []).append(value)
    total_required = len(groups)
    for group, candidates in groups.items():
        match = False
        for cand in candidates:
            if cand and cand in a:
                match = True
                break
        if match:
            found += 1

    # accuracy is fraction found / total_required. If nothing required (no trip details), return 0.0
    if total_required == 0:
        accuracy = 0.0
    else:
        accuracy = found / total_required

    accuracy = float(min(1.0, max(0.0, accuracy)))

    return relevance, accuracy

# -----------------------------
# Sidebar Metrics
# -----------------------------
def render_sidebar():
    st.sidebar.header("📊 Evaluation Metrics (Live)")
    st.sidebar.metric("Relevance", f"{st.session_state.relevance:.2f}/1.0")
    st.sidebar.progress(min(st.session_state.relevance, 1.0))
    st.sidebar.metric("Info Accuracy", f"{st.session_state.accuracy:.2f}/1.0")
    st.sidebar.progress(min(st.session_state.accuracy, 1.0))

render_sidebar()

# -----------------------------
# Step 1 — Ask Trip Details
# -----------------------------
if not st.session_state.trip_details_submitted:
    st.header("✈️ Enter Your Trip Details")

    with st.form(key="trip_form"):
        st.session_state.departure_city = st.text_input("Departure City", value=st.session_state.departure_city)
        st.session_state.destination_city = st.text_input("Destination City", value=st.session_state.destination_city)
        st.session_state.departure_date = st.date_input(
            "Departure Date", value=st.session_state.departure_date or datetime.date.today(), min_value=datetime.date.today()
        )
        st.session_state.return_date = st.date_input(
            "Return Date", value=st.session_state.return_date or datetime.date.today(), min_value=datetime.date.today()
        )
        trip_submit = st.form_submit_button("Continue")

    if trip_submit:
        if not st.session_state.departure_city or not st.session_state.destination_city:
            st.warning("Please fill in both departure and destination cities.")
        elif st.session_state.return_date < st.session_state.departure_date:
            st.warning("Return date cannot be before departure date.")
        else:
            st.session_state.trip_details_submitted = True
            st.success("Trip details saved! Generating your itinerary...")
            st.rerun()

# -----------------------------
# Step 2 — Generate Original Trip Plan
# -----------------------------
if st.session_state.trip_details_submitted:
    st.header("Your AI-Generated Trip Plan ✈️")
    st.markdown(
        f"**From:** {st.session_state.departure_city} → **To:** {st.session_state.destination_city}<br>"
        f"**Travel Dates:** {st.session_state.departure_date} → {st.session_state.return_date}",
        unsafe_allow_html=True,
    )

    # Generate original plan once
    if st.session_state.last_response == "":
        with st.spinner("Building your personalized travel itinerary..."):
            auto_query = (
                f"Create a complete travel itinerary for a trip from "
                f"{st.session_state.departure_city} to {st.session_state.destination_city}, "
                f"departing on {st.session_state.departure_date} and returning on {st.session_state.return_date}. "
                f"Include sightseeing, food, restarunt, budgeting, travel tips and day-wise break-up, Shops at airport/Railway Station, Most usefull Mobile Apps at that location with download links ."
            )
            payload = {"question": auto_query, "chat_history": ""}

            try:
                response = requests.post(f"{BASE_URL}/query", json=payload, timeout=120)
            except Exception as e:
                st.error(f"Failed to reach backend: {e}")
                response = None

        if response and response.status_code == 200:
            result = response.json()
            backend_answer = result.get("answer", "").strip() if isinstance(result.get("answer", ""), str) else str(result.get("answer", ""))
            # token accounting — prefer backend values, fallback to estimator
            st.session_state.tokens_in = int(result.get("tokens_in") or estimate_tokens(auto_query))
            st.session_state.tokens_out = int(result.get("tokens_out") or estimate_tokens(backend_answer))
            st.session_state.total_tokens = int(result.get("total_tokens") or st.session_state.tokens_in + st.session_state.tokens_out)

            # store answer and update UI
            st.session_state.last_response = backend_answer
            st.session_state.plans.append(("Original Plan", backend_answer))

            # compute evaluation locally
            try:
                rel, acc = compute_evaluation(
                    user_query=auto_query,
                    answer_text=backend_answer,
                    departure_city=st.session_state.departure_city,
                    destination_city=st.session_state.destination_city,
                    departure_date=st.session_state.departure_date,
                    return_date=st.session_state.return_date,
                )
            except Exception:
                # in unlikely error, fallback to conservative defaults (not hard-coded nice values)
                rel, acc = 0.5, 0.5

            # update session state
            st.session_state.relevance = rel
            st.session_state.accuracy = acc

            # re-render sidebar & page
            render_sidebar()
            st.rerun()
        else:
            st.error("Backend failed to respond.")


        # Display all existing plans
    for title, plan in st.session_state.plans:
        st.markdown(
            f"<h4>📝 {title}</h4>"
            f"<div style='padding:10px; border:1px solid #DDD; border-radius:10px;'>{plan}</div>",
            unsafe_allow_html=True,
        )


    # -----------------------------
    # ⭐ Refine or Modify Your Trip Plan
    # -----------------------------
    st.subheader("✨ Refine or Modify Your Trip Plan")

    suggestion_map = {
        "🏔 More Adventure": "Make the itinerary more adventurous with hiking, trekking, and outdoor activities.",
        "🍽 More Food Spots": "Add more food recommendations, street foods, restaurants, and local cuisine.",
        "💸 Budget-Friendly": "Modify the trip to be more budget-friendly with cheaper stays and low-cost activities.",
        "🧘 Slow & Relaxed": "Make the trip more relaxed with slow travel and less hectic days.",
        "👨‍👩‍👧 Family-Friendly": "Modify the trip to be family-friendly with kid-safe attractions."
    }

    selected_option = st.radio(
        "Choose a suggestion:",
        list(suggestion_map.keys()),
        index=None
    )

    modify_option = suggestion_map.get(selected_option) if selected_option else None

    # Custom user input box
    user_mod_query = st.text_input(
        "Or enter your own suggestion to refine the trip:",
        placeholder="Example: Add more beaches, make it luxury, reduce travel time, include shopping..."
    )

    if st.button("Update Trip Plan"):
        final_mod = modify_option or (user_mod_query.strip() if user_mod_query.strip() else None)
        if final_mod:
            with st.spinner("Generating modified itinerary..."):
                refine_query = (
                    f"Refine this itinerary based on user request: {final_mod}. "
                    f"Original & previous itineraries: {st.session_state.plans}"
                )
                payload = {"question": refine_query, "chat_history": st.session_state.last_response}
                try:
                    response = requests.post(f"{BASE_URL}/query", json=payload, timeout=120)
                except Exception as e:
                    st.error(f"Failed to reach backend: {e}")
                    response = None

                if response and response.status_code == 200:
                    result = response.json()
                    new_plan = result.get("answer", "").strip() if isinstance(result.get("answer", ""), str) else str(result.get("answer", ""))
                    st.session_state.plans.append((f"Refined Plan ({final_mod})", new_plan))

                    # Token handling
                    tokens_in_new = int(result.get("tokens_in") or estimate_tokens(refine_query))
                    tokens_out_new = int(result.get("tokens_out") or estimate_tokens(new_plan))
                    st.session_state.tokens_in += tokens_in_new
                    st.session_state.tokens_out += tokens_out_new
                    st.session_state.total_tokens = int(result.get("total_tokens") or st.session_state.tokens_in + st.session_state.tokens_out)

                    # compute evaluation AGAIN — use refine_query as the "user query"
                    try:
                        rel_new, acc_new = compute_evaluation(
                            user_query=refine_query,
                            answer_text=new_plan,
                            departure_city=st.session_state.departure_city,
                            destination_city=st.session_state.destination_city,
                            departure_date=st.session_state.departure_date,
                            return_date=st.session_state.return_date,
                        )
                    except Exception:
                        rel_new, acc_new = st.session_state.relevance, st.session_state.accuracy

                    # Update scores: we choose to average previous and new to give a running estimate
                    st.session_state.relevance = float((st.session_state.relevance + rel_new) / 2.0)
                    st.session_state.accuracy = float((st.session_state.accuracy + acc_new) / 2.0)

                    render_sidebar()
                    st.rerun()
                else:
                    st.error("Backend failed to refine itinerary.")
        else:
            st.warning("Please choose or type a modification to apply.")
    # ----------------------------------------------------

    # -----------------------------
    # 🛡 Travel Insurance Suggestions (Squaremouth API)
    # -----------------------------
    st.subheader("🛡 Travel Insurance Options (Recommended)")
    st.markdown("Get insurance options based on your trip details. Powered by your backend insurance API.")

    # Pre-fill using trip details already collected
    ins_destination = st.session_state.destination_city
    ins_start = st.session_state.departure_date
    ins_end = st.session_state.return_date

    ins_age = st.number_input("Traveler Age", min_value=1, max_value=90, value=10)
    ins_coverage = st.selectbox(
        "Coverage Level",
        ["standard", "comprehensive", "budget"]
    )

    if st.button("Get Travel Insurance Quotes"):
        if not ins_destination or not ins_start or not ins_end:
            st.warning("Trip details missing. Please ensure destination and dates are filled.")
        else:
            with st.spinner("Fetching insurance options..."):
                payload = {
                    "destination": ins_destination,
                    "trip_start": ins_start.strftime("%Y-%m-%d"),
                    "trip_end": ins_end.strftime("%Y-%m-%d"),
                    "age": ins_age,
                    "travelers": 1,
                    "coverage_type": ins_coverage
                }

                try:
                    resp = requests.post(
                        f"{BASE_URL}/insurance/squaremouth",
                        json=payload,
                        timeout=60
                    )

                    if resp.status_code == 200:
                        data = resp.json()
                        quotes = data.get("insurance_quotes", {}).get("plans", [])

                        if quotes:
                            # st.success("Insurance plans found:")
                            for p in quotes:
                                with st.expander(f"{p.get('provider')} — {p.get('plan_name')}"):
                                    st.write(f"**Price:** ₹{p.get('price'):,}")
                                    st.write(f"**Medical Coverage:** {p.get('medical')}")
                                    st.write(f"**Trip Cancellation:** {p.get('trip_cancellation')}")
                                    st.write(f"**Baggage Loss:** {p.get('baggage')}")
                                    st.write(f"**Evacuation:** {p.get('evacuation')}")
                        else:
                            st.info("No plans available for your trip details.")

                    else:
                        st.error(f"Insurance API error: {resp.text}")

                except Exception as e:
                    st.error(f"Failed to connect to insurance API: {e}")


    st.markdown("---")


    # Token Count Display (moved from sidebar to bottom)
    # ----------------------------------------------------
    st.markdown("---")
    st.subheader("Token Usage Summary")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.metric("Tokens In", st.session_state.tokens_in)

    with col2:
        st.metric("Tokens Out", st.session_state.tokens_out)

    with col3:
        st.metric("Total Tokens", st.session_state.total_tokens)


    def generate_word_doc(all_plans, relevance, accuracy):
        doc = Document()
        doc.add_heading("AI Trip Plan Summary", level=1)
        doc.add_paragraph(f"Date: {datetime.date.today().strftime('%d-%b-%Y')}")
        doc.add_paragraph(f"From: {st.session_state.departure_city}")
        doc.add_paragraph(f"To: {st.session_state.destination_city}")
        doc.add_paragraph(f"Travel Dates: {st.session_state.departure_date} → {st.session_state.return_date}")
        for title, plan in all_plans:
            doc.add_heading(title, level=2)
            # preserve basic line breaks
            for line in plan.splitlines():
                doc.add_paragraph(line)
        doc.add_paragraph("\nEvaluation Metrics:")
        doc.add_paragraph(f"Relevance: {relevance:.2f}/1.0")
        doc.add_paragraph(f"Info Accuracy: {accuracy:.2f}/1.0")
        doc.add_paragraph(f"Input Tokens: {st.session_state.tokens_in}")
        doc.add_paragraph(f"Output Tokens: {st.session_state.tokens_out}")
        doc.add_paragraph(f"Total Tokens: {st.session_state.total_tokens}")
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    word_buffer = generate_word_doc(
        st.session_state.plans,
        st.session_state.relevance,
        st.session_state.accuracy,
    )

    # -----------------------------
    # Export all plans to Word
    # -----------------------------
    st.divider()
    st.subheader("📄 Export All Trip Plans (Original + All Modifications)")

    st.download_button(
        label="💾 Download ALL Itineraries as Word File",
        data=word_buffer,
        file_name=f"Trip_Plan_{datetime.date.today().isoformat()}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
